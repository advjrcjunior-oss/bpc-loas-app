"""
Servico de consulta CPF/CNPJ via cpfcnpj.com.br
Pacote F: retorna nome, nascimento, endereco completo, telefone.
Token: CPFCNPJ_TOKEN no .env
"""
import os
import re
import requests

CPFCNPJ_TOKEN = os.environ.get("CPFCNPJ_TOKEN", "")
BASE_URL = "https://api.cpfcnpj.com.br"


def limpar_cpf(cpf: str) -> str:
    return re.sub(r"\D", "", cpf)


def consultar_cpf(cpf: str) -> dict:
    """
    Consulta CPF e retorna dados cadastrais.
    Retorna dict com: nome, nascimento, endereco{logradouro, numero, complemento, bairro, cidade, uf, cep}, telefone
    Raises ValueError se CPF invalido ou API retornar erro.
    """
    cpf_limpo = limpar_cpf(cpf)
    if len(cpf_limpo) != 11:
        raise ValueError(f"CPF invalido: {cpf}")

    token = CPFCNPJ_TOKEN
    if not token:
        raise ValueError("CPFCNPJ_TOKEN nao configurado no .env")

    # Pacote F: dados completos (nome, nascimento, endereco, telefone)
    url = f"{BASE_URL}/{token}/6/json/{cpf_limpo}"
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()
    data = resp.json()

    if data.get("codigo_erro"):
        raise ValueError(f"Erro API CPFCNPJ: {data.get('mensagem_erro', 'desconhecido')}")

    return {
        "cpf": cpf_limpo,
        "nome": data.get("nome", "").strip(),
        "nascimento": data.get("nascimento", "").strip(),
        "sexo": data.get("sexo", "").strip(),
        "mae": data.get("mae", "").strip(),
        "endereco": {
            "logradouro": data.get("logradouro", "").strip(),
            "numero": data.get("numero", "").strip(),
            "complemento": data.get("complemento", "").strip(),
            "bairro": data.get("bairro", "").strip(),
            "cidade": data.get("cidade", "").strip(),
            "uf": data.get("uf", "").strip(),
            "cep": re.sub(r"\D", "", data.get("cep", "")),
        },
        "telefone": data.get("telefone", "").strip(),
        "raw": data,
    }


def consultar_cnpj(cnpj: str) -> dict:
    """Consulta CNPJ e retorna dados cadastrais."""
    cnpj_limpo = re.sub(r"\D", "", cnpj)
    if len(cnpj_limpo) != 14:
        raise ValueError(f"CNPJ invalido: {cnpj}")

    token = CPFCNPJ_TOKEN
    if not token:
        raise ValueError("CPFCNPJ_TOKEN nao configurado no .env")

    url = f"{BASE_URL}/{token}/6/json/{cnpj_limpo}"
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()
    data = resp.json()

    if data.get("codigo_erro"):
        raise ValueError(f"Erro API CPFCNPJ: {data.get('mensagem_erro', 'desconhecido')}")

    return data


def validar_dados_cliente(cpf: str, nome_esperado: str = None) -> dict:
    """
    Consulta CPF e valida contra dados esperados.
    Retorna dict com dados + campo 'validacao' indicando discrepancias.
    """
    dados = consultar_cpf(cpf)
    validacao = {"cpf_valido": True, "alertas": []}

    if nome_esperado:
        nome_api = dados["nome"].upper()
        nome_esp = nome_esperado.upper()
        if nome_esp not in nome_api and nome_api not in nome_esp:
            validacao["alertas"].append(
                f"Nome diverge: API='{dados['nome']}' vs Esperado='{nome_esperado}'"
            )

    if not dados["endereco"]["cep"]:
        validacao["alertas"].append("Endereco sem CEP na base CPFCNPJ")

    dados["validacao"] = validacao
    return dados


def _formatar_cpf(cpf: str) -> str:
    c = limpar_cpf(cpf)
    if len(c) == 11:
        return f"{c[:3]}.{c[3:6]}.{c[6:9]}-{c[9:]}"
    return cpf


def gerar_declaracao_residencia(
    nome: str,
    cpf: str,
    logradouro: str,
    numero: str,
    bairro: str,
    cidade: str,
    uf: str,
    cep: str,
    output_dir: str = None,
) -> str:
    """
    Gera Declaracao de Residencia em DOCX (Lei 7.115/83).
    Retorna caminho do arquivo gerado.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Titulo
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("DECLARAÇÃO DE RESIDÊNCIA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"

    # Subtitulo legal
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("(Lei nº 7.115, de 29 de agosto de 1983)")
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.italic = True

    doc.add_paragraph()  # espaco

    # Corpo
    cpf_fmt = _formatar_cpf(cpf)
    cep_fmt = f"{cep[:5]}-{cep[5:]}" if len(re.sub(r'\\D', '', cep)) == 8 else cep
    endereco_completo = f"{logradouro}, nº {numero}"
    if bairro:
        endereco_completo += f", {bairro}"
    endereco_completo += f", {cidade}/{uf}"
    if cep:
        endereco_completo += f", CEP {cep_fmt}"

    meses = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
    ]
    hoje = datetime.now()
    data_extenso = f"{hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"

    texto = (
        f"Eu, {nome.upper()}, portador(a) do CPF nº {cpf_fmt}, "
        f"DECLARO, para os devidos fins de direito e sob as penas da lei, "
        f"que RESIDO no endereço: {endereco_completo}."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.font.size = Pt(12)
    run.font.name = "Arial"
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.first_line_indent = Cm(1.5)

    doc.add_paragraph()  # espaco

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p2.add_run(
        "Declaro, ainda, que estou ciente de que a falsidade desta declaração "
        "pode implicar em sanção penal prevista no art. 299 do Código Penal, "
        "além de sanções civis e administrativas."
    )
    run.font.size = Pt(12)
    run.font.name = "Arial"
    p2.paragraph_format.line_spacing = Pt(18)
    p2.paragraph_format.first_line_indent = Cm(1.5)

    doc.add_paragraph()  # espaco

    # Data e local
    local = doc.add_paragraph()
    local.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = local.add_run(f"{cidade}/{uf}, {data_extenso}.")
    run.font.size = Pt(12)
    run.font.name = "Arial"

    doc.add_paragraph()  # espaco
    doc.add_paragraph()  # espaco

    # Assinatura
    assinatura = doc.add_paragraph()
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = assinatura.add_run("_" * 50)
    run.font.size = Pt(12)
    run.font.name = "Arial"

    nome_p = doc.add_paragraph()
    nome_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = nome_p.add_run(nome.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"

    cpf_p = doc.add_paragraph()
    cpf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cpf_p.add_run(f"CPF: {cpf_fmt}")
    run.font.size = Pt(11)
    run.font.name = "Arial"

    # Salvar
    if not output_dir:
        output_dir = os.path.dirname(os.path.abspath(__file__))
    filename = f"declaracao_residencia_{limpar_cpf(cpf)}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Uso: python cpfcnpj_service.py <CPF>")
        sys.exit(1)
    resultado = consultar_cpf(sys.argv[1])
    import json
    print(json.dumps(resultado, indent=2, ensure_ascii=False))
